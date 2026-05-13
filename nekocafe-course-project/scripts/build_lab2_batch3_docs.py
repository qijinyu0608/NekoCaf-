from __future__ import annotations

from pathlib import Path

from docx import Document

from build_lab2_docs import add_cover, add_table_caption, apply_page_style
from lab2_batch3_common import ADR_RECORDS, ROADMAP, adr_markdown_dir, adr_package_dir, roadmap_docx_path
from lab2_common import META, REFERENCES, ensure_dirs


def write_markdown(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def build_adr_markdown(record: dict) -> str:
    alternatives = "\n".join(f"- {item}" for item in record["alternatives"])
    consequences = "\n".join(f"- {item}" for item in record["consequences"])
    refs = "\n".join(f"- {item}" for item in REFERENCES[:3])
    return f"""# {record['id']} {record['title']}

- Status: {record['status']}
- Topic: {record['topic']}

## Context
{record['context']}

## Decision
{record['decision']}

## Consequences
{consequences}

## Alternatives Considered
{alternatives}

## References
{refs}
"""


def build_adr_docx(record: dict, path: Path) -> None:
    doc = Document()
    apply_page_style(doc)
    add_cover(doc, f"ADR {record['id']} {record['title']}", f"D2-7 Architecture Decision Record {record['id']}", "D2-7")

    doc.add_heading("1 Status", level=1)
    doc.add_paragraph(record["status"])
    doc.add_heading("2 Topic", level=1)
    doc.add_paragraph(record["topic"])
    doc.add_heading("3 Context", level=1)
    doc.add_paragraph(record["context"])
    doc.add_heading("4 Decision", level=1)
    doc.add_paragraph(record["decision"])
    doc.add_heading("5 Consequences", level=1)
    for item in record["consequences"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("6 Alternatives Considered", level=1)
    for item in record["alternatives"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("7 References", level=1)
    for item in REFERENCES[:3]:
        doc.add_paragraph(item)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_adr_readme_docx(path: Path) -> None:
    doc = Document()
    apply_page_style(doc)
    add_cover(doc, "D2-7 架构决策记录 ADR 集 README", "D2-7 ADR Package README", "D2-7")

    doc.add_heading("1 本目录用途", level=1)
    doc.add_paragraph("本目录存放实验二的 ADR 集。README 用于说明编号规则、主决策主题，以及 Markdown 托管目录与实验交付目录之间的关系。")

    doc.add_heading("2 编号与同步规则", level=1)
    bullets = [
        "0001 固定用于确立 ADR 制度。",
        "0002 起按决策出现顺序递增，不重排历史编号。",
        "仓库根目录 docs/adr/lab2/ 下的 Markdown 为托管主版本，D2-7 文件夹中的 docx 为课程交付版。",
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("3 ADR 清单", level=1)
    add_table_caption(doc, "表 3-1 ADR 编号与主题")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, header in enumerate(["编号", "主题", "标题"]):
        table.rows[0].cells[idx].text = header
    for record in ADR_RECORDS:
        row = table.add_row().cells
        row[0].text = record["id"]
        row[1].text = record["topic"]
        row[2].text = record["title"]

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_adr_outputs() -> None:
    package_dir = adr_package_dir()
    markdown_dir = adr_markdown_dir()
    package_dir.mkdir(parents=True, exist_ok=True)
    markdown_dir.mkdir(parents=True, exist_ok=True)

    readme_md = [
        "# Lab2 ADR Index",
        "",
        "本目录为实验二 ADR 的 Git 托管主目录，课程交付版 docx 位于 `docs/lab2/source/...ADRSet_v1.0/`。",
        "",
    ]
    for record in ADR_RECORDS:
        md_name = f"{record['id']}-{record['topic'].replace(' ', '-')}.md"
        readme_md.append(f"- [{record['id']} {record['topic']}]({md_name})")
        write_markdown(markdown_dir / md_name, build_adr_markdown(record))
        build_adr_docx(record, package_dir / f"{record['id']}-{record['topic']}.docx")
    write_markdown(markdown_dir / "README.md", "\n".join(readme_md) + "\n")
    build_adr_readme_docx(package_dir / "README.docx")


def build_roadmap_docx() -> None:
    doc = Document()
    apply_page_style(doc)
    add_cover(doc, "架构演进路线图（12 个月）", "D2-8 Architecture Roadmap", "D2-8")

    doc.add_heading("1 现状评估", level=1)
    for item in ROADMAP["current_state"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2 目标态描述", level=1)
    for item in ROADMAP["target_state"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3 阶段规划", level=1)
    for item in ROADMAP["milestones"]:
        doc.add_paragraph(f"{item['code']}（{item['months']}）：{item['title']}", style="Heading 2")
        doc.add_paragraph(item["focus"])
        doc.add_paragraph("关键交付物：", style="Heading 3")
        for deliverable in item["deliverables"]:
            doc.add_paragraph(deliverable, style="List Bullet")
        doc.add_paragraph("成功的样子（Definition of Success）：", style="Heading 3")
        doc.add_paragraph(item["definition_of_success"])

    doc.add_heading("4 风险登记册", level=1)
    add_table_caption(doc, "表 4-1 风险登记册")
    risk_table = doc.add_table(rows=1, cols=3)
    risk_table.style = "Table Grid"
    for idx, header in enumerate(["风险", "影响", "缓解策略"]):
        risk_table.rows[0].cells[idx].text = header
    for risk, impact, mitigation in ROADMAP["risk_register"]:
        row = risk_table.add_row().cells
        row[0].text = risk
        row[1].text = impact
        row[2].text = mitigation

    doc.add_heading("5 退出与回滚策略", level=1)
    for item in ROADMAP["rollback"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6 参考文献", level=1)
    for item in REFERENCES[:4]:
        doc.add_paragraph(item)

    roadmap_docx_path().parent.mkdir(parents=True, exist_ok=True)
    doc.save(roadmap_docx_path())


if __name__ == "__main__":
    ensure_dirs()
    build_adr_outputs()
    build_roadmap_docx()
