from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from docx_style_policy import apply_black_text_policy


ROOT = Path(__file__).resolve().parents[1]
FOUNDATION_DIR = ROOT / "docs" / "foundation"


def apply_page_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    apply_black_text_policy(
        doc,
        heading_sizes={
            "Title": 18,
            "Heading 1": 16,
            "Heading 2": 14,
        },
    )

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.text = "《软件工程》实验报告 — NekoCafé 项目"

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "NekoCafé 基座文档"


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    p.style = doc.styles["Title"]

    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_project_overview() -> None:
    doc = Document()
    apply_page_style(doc)
    add_title(doc, "NekoCafé 统一项目概述", "课程四次实验共享基线文档")

    doc.add_heading("1. 项目目标", level=1)
    for item in [
        "将四次软件工程实验组织成一个连续演进的项目，而不是四份互相割裂的作业。",
        "用共享的需求、术语、追溯矩阵和命名规则，减少后续实验返工。",
        "为实验三与实验四提前准备统一的 Python monorepo 骨架。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. 案例摘要", level=1)
    doc.add_paragraph(
        "案例统一采用 NekoCafé 猫咪主题餐饮预约平台，覆盖顾客预约、会员管理、店员调度、"
        "猫咪健康档案和运营分析等业务。"
    )

    doc.add_heading("3. 干系人与关注点", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["角色", "主要目标", "对应实验关注"]
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = value
    rows = [
        ("顾客", "预约便捷、推荐准确、信息透明", "实验一需求获取、实验四用户旅程测试"),
        ("店员", "操作高效、调度顺畅、异常可见", "实验一需求、实验二服务设计"),
        ("运营总监", "跨门店分析、活动与合规", "实验二架构、实验三观测与交付"),
        ("QA/运维", "质量门禁、回滚、定位效率", "实验三流水线、实验四质量工程"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    doc.add_heading("4. 统一命名规则", level=1)
    naming_items = [
        "功能需求使用 FR-001 起编号，非功能需求使用 NFR-001 起编号。",
        "用例 ID 使用 UC- 前缀，限界上下文使用 BC- 前缀，服务使用 SVC- 前缀。",
        "自动化测试用例使用 TC- 前缀，并在实验四回填到 RTM 扩展矩阵。",
        "所有术语以 glossary-rtm-master.xlsx 为唯一主源，不在单个实验内重新命名。",
    ]
    for item in naming_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5. 目录约束", level=1)
    doc.add_paragraph(
        "工作区将文档基线、实验产物、工程代码和模板副本分离存放。老师原始任务书与模板保留在外层目录，"
        "不在统一工作区内直接覆盖。"
    )

    doc.add_heading("6. 服务范围基线", level=1)
    for item in [
        "可运行范围：reservation-service、member-service。",
        "设计表达范围：order、store-ops、cat-health、recommendation。",
        "实验二至少表达 6 个限界上下文，但实验三和实验四只对两个核心服务做 PoC。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(FOUNDATION_DIR / "project-overview.docx")


def build_artifact_map() -> None:
    doc = Document()
    apply_page_style(doc)
    add_title(doc, "共享基线到实验交付物映射表", "Foundation → D1 / D2 / D3 / D4")

    doc.add_heading("1. 设计原则", level=1)
    for item in [
        "共享基线只维护一份主数据，正式交付物按实验拆分表达。",
        "实验一产出需求与术语主表，实验二到实验四全部回引这份主表。",
        "源文件、导出件和提交件分目录保存，避免同一文档多处散落。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. 映射总表", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["共享基线", "主要内容", "映射到实验交付", "使用方式"]
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = value

    mappings = [
        (
            "project-overview.docx",
            "案例、角色、命名、目录规范",
            "D1-1 / D1-3 / D2-1 / D3-1 / D4-1",
            "作为背景和术语一致性的引用源",
        ),
        (
            "requirements-baseline.xlsx",
            "FR / NFR 主表、来源、优先级、验收准则",
            "D1-2 / D1-3 / D2-1 / D4-2",
            "实验一形成基线，后续按视图过滤与补充",
        ),
        (
            "glossary-rtm-master.xlsx",
            "Glossary、RTM、命名规则",
            "D1-5 / D4-8 / D2 追溯说明",
            "同一主表扩展列，不复制第二份",
        ),
        (
            "project/ 代码骨架",
            "服务入口、共享常量、测试目录、infra 占位",
            "D3-2 / D4-3 / D4-7",
            "实验三实现 PoC，实验四追加质量体系",
        ),
    ]
    for row in mappings:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    doc.add_heading("3. 执行节奏", level=1)
    ordered = [
        "先完善 foundation 主表，再写实验一正式文档。",
        "实验二所有图、OpenAPI、ADR 必须使用实验一的需求 ID 与术语。",
        "实验三从 project/ 目录接手，实现两个服务的最小运行链路。",
        "实验四直接在同一项目上补测试、质量看板、AI 评审与缺陷追踪。",
    ]
    for item in ordered:
        doc.add_paragraph(item, style="List Number")

    doc.save(FOUNDATION_DIR / "artifact-map.docx")


if __name__ == "__main__":
    FOUNDATION_DIR.mkdir(parents=True, exist_ok=True)
    build_project_overview()
    build_artifact_map()
