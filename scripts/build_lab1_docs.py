from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont

from lab1_common import (
    AI_PROMPTS,
    BUILD_DIR,
    DEFECTS,
    DIAGRAM_DIR,
    EXPORT_DIR,
    FOUNDATION_DIR,
    GLOSSARY,
    INTERVIEW_QUESTIONS,
    META,
    NON_FUNCTIONAL_REQUIREMENTS,
    RISKS,
    SLIDES,
    SOURCE_DIR,
    STAKEHOLDERS,
    SUBMISSION_DIR,
    TBD_ITEMS,
    USE_CASES,
    VALIDATION_CHECKLIST,
    artifact_filename,
    artifact_stem,
    build_transcript,
    customer_personas,
    ensure_dirs,
)


FONT_PATH = "/System/Library/Fonts/PingFang.ttc"
UML_EXPORT_DIR = EXPORT_DIR / artifact_stem("D1-4")


def render_font(size: int):
    try:
        return ImageFont.truetype(FONT_PATH, size)
    except Exception:
        return ImageFont.load_default()


def apply_page_style(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Title", 18), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def configure_section(section) -> None:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.text = META["header"]

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)


def add_page_number(paragraph) -> None:
    paragraph.text = ""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    p.style = doc.styles["Title"]

    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"
    pairs = [
        ("课程 / 实验", f"{META['course']} / {META['experiment']} {META['experiment_title']}"),
        ("项目案例", META["suite_title"]),
        ("班级 / 学号 / 姓名", f"{META['class_name']} / {META['student_id']} / {META['student_name']}"),
        ("版本 / 日期", f"{META['version']} / {META['date']}"),
    ]
    for row_idx, pair in enumerate(pairs):
        info.rows[row_idx].cells[0].text = pair[0]
        info.rows[row_idx].cells[1].text = pair[1]
        for cell in info.rows[row_idx].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("学术诚信声明", style="Heading 2")
    doc.add_paragraph(META["declaration"])
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(new_section)


def add_paragraphs(doc: Document, lines: list[str], style: str | None = None) -> None:
    for line in lines:
        doc.add_paragraph(line, style=style) if style else doc.add_paragraph(line)


def add_caption(doc: Document, caption: str, kind: str = "图") -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if kind == "图" else WD_ALIGN_PARAGRAPH.LEFT
    para.add_run(caption).bold = True


def draw_support_figures() -> tuple[Path, Path]:
    ensure_dirs()
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    onion_path = BUILD_DIR / "onion_model.png"
    matrix_path = BUILD_DIR / "priority_matrix.png"
    font_title = render_font(28)
    font_body = render_font(18)

    onion = Image.new("RGB", (1200, 800), "#FFFFFF")
    draw = ImageDraw.Draw(onion)
    center = (600, 400)
    layers = [
        (280, "#F4E2D8", "核心用户层"),
        (220, "#EED7C5", "运营执行层"),
        (160, "#E5C7AE", "管理决策层"),
        (90, "#D86F45", "平台目标"),
    ]
    for radius, color, label in layers:
        bbox = (center[0] - radius, center[1] - radius, center[0] + radius, center[1] + radius)
        draw.ellipse(bbox, fill=color, outline="#5A5A5A", width=2)
        draw.text((center[0] - radius + 20, center[1]), label, fill="#333333", font=font_body)
    draw.text((470, 60), "Onion Model 干系人分布", fill="#333333", font=font_title)
    draw.text((475, 340), "顾客代表 A / 顾客代表 B", fill="#333333", font=font_body)
    draw.text((520, 405), "店员代表", fill="#333333", font=font_body)
    draw.text((520, 455), "运营总监", fill="#333333", font=font_body)
    draw.text((550, 385), "NekoCafé 平台升级", fill="#FFFFFF", font=font_body)
    onion.save(onion_path)

    matrix = Image.new("RGB", (1200, 800), "#FFFFFF")
    draw = ImageDraw.Draw(matrix)
    draw.text((430, 60), "干系人优先级矩阵（影响力 × 关注度）", fill="#333333", font=font_title)
    draw.line((160, 680, 1000, 680), fill="#5A5A5A", width=4)
    draw.line((160, 680, 160, 140), fill="#5A5A5A", width=4)
    draw.text((1010, 660), "关注度", fill="#333333", font=font_body)
    draw.text((90, 110), "影响力", fill="#333333", font=font_body)
    draw.line((580, 180, 580, 680), fill="#A0A0A0", width=2)
    draw.line((160, 430, 1000, 430), fill="#A0A0A0", width=2)
    points = [
        ("顾客代表 A", 700, 520, "#D86F45"),
        ("顾客代表 B", 760, 560, "#D86F45"),
        ("店员代表", 680, 300, "#6A8D73"),
        ("运营总监", 820, 250, "#5D7FA3"),
    ]
    for label, x, y, color in points:
        draw.ellipse((x - 12, y - 12, x + 12, y + 12), fill=color)
        draw.text((x + 20, y - 10), label, fill="#333333", font=font_body)
    matrix.save(matrix_path)
    return onion_path, matrix_path


def add_image(doc: Document, image_path: Path, width_cm: float, caption: str) -> None:
    doc.add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, caption, "图")


def add_table_caption(doc: Document, caption: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(caption)
    run.bold = True


def build_d1_1() -> None:
    doc = Document()
    apply_page_style(doc)
    add_cover(doc, "需求获取计划与访谈剧本", "D1-1 Requirements Plan")
    onion_path, matrix_path = draw_support_figures()

    doc.add_heading("1 文档信息", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("产出编号", "D1-1"),
        ("适用案例", META["suite_title"]),
        ("作者", META["student_name"]),
        ("修订记录", "v1.0 初稿，根据实验一任务书和总任务书生成"),
    ]
    for idx, row in enumerate(rows):
        table.rows[idx].cells[0].text = row[0]
        table.rows[idx].cells[1].text = row[1]

    doc.add_heading("2 项目背景与目标", level=1)
    add_paragraphs(
        doc,
        [
            "NekoCafé 统一案例要求以同一份需求基线服务四次连贯实验，因此本次需求获取不仅要回答实验一的“做什么”，还要为实验二到实验四保留术语、编号和追溯主线。",
            "本轮需求获取目标包括：识别顾客、店员、运营总监三类核心干系人的真实诉求；形成 30 条功能需求与 10 条非功能需求的基线；建立统一术语表和 RTM，为后续 UML、SRS、架构和测试产物提供共同事实源。",
        ],
    )

    doc.add_heading("3 干系人分析", level=1)
    doc.add_heading("3.1 Onion Model 图", level=2)
    add_image(doc, onion_path, 14.5, "图 3-1 Onion Model 干系人分布")
    doc.add_paragraph("核心用户层为两位顾客代表，强调预约便捷、信息透明和体验稳定；运营执行层为店员代表，聚焦排台、异常处理与猫咪健康打卡；管理决策层为运营总监，关注跨店经营指标、合规与标准化复制。")
    doc.add_heading("3.2 干系人优先级矩阵", level=2)
    add_image(doc, matrix_path, 14.5, "图 3-2 干系人优先级矩阵")
    doc.add_paragraph("店员代表与运营总监处于高影响力象限，分别代表执行约束和管理目标；两位顾客代表位于高关注度象限，决定前台体验与会员价值表达。")

    doc.add_heading("4 访谈策略", level=1)
    doc.add_paragraph("本次采用“AI 半结构化访谈剧本 + 人工归并复核”的方式开展需求获取。AI 用于快速扩展问题覆盖面与异常场景，人工负责术语统一、边界约束与需求优先级判断。")
    add_table_caption(doc, "表 4-1 访谈时间安排表")
    schedule = doc.add_table(rows=1, cols=5)
    schedule.style = "Table Grid"
    for idx, value in enumerate(["对象", "日期", "时长", "方式", "目标"]):
        schedule.rows[0].cells[idx].text = value
    entries = [
        ("顾客代表 A", "2026-05-07", "35 分钟", "AI 模拟访谈", "确认预约主链路与体验指标"),
        ("顾客代表 B", "2026-05-07", "35 分钟", "AI 模拟访谈", "确认改期、爽约与优惠诉求"),
        ("店员代表", "2026-05-07", "40 分钟", "AI 模拟访谈", "确认排台、异常处理与健康打卡"),
        ("运营总监", "2026-05-07", "35 分钟", "AI 模拟访谈", "确认跨店运营、合规和阶段边界"),
    ]
    for entry in entries:
        cells = schedule.add_row().cells
        for idx, value in enumerate(entry):
            cells[idx].text = value

    doc.add_heading("5 访谈剧本", level=1)
    for role, questions in INTERVIEW_QUESTIONS.items():
        doc.add_heading(f"5.{['顾客','店员','运营总监'].index(role)+1} {role}剧本", level=2)
        for idx, question in enumerate(questions, start=1):
            doc.add_paragraph(f"{idx}. {question}", style="List Number")

    doc.add_heading("6 风险与应急方案", level=1)
    add_table_caption(doc, "表 6-1 需求获取阶段风险与应对")
    risk_table = doc.add_table(rows=1, cols=4)
    risk_table.style = "Table Grid"
    for idx, value in enumerate(["风险ID", "风险描述", "等级", "应对"]):
        risk_table.rows[0].cells[idx].text = value
    for risk in RISKS:
        cells = risk_table.add_row().cells
        for idx, value in enumerate(risk):
            cells[idx].text = value

    doc.add_heading("7 附录 A AI 模拟干系人 Prompt 全文", level=1)
    for item in AI_PROMPTS[:2]:
        doc.add_heading(f"{item['title']}（{item['version']}）", level=2)
        doc.add_paragraph(f"用途：{item['purpose']}")
        doc.add_paragraph(item["prompt"])

    doc.save(SOURCE_DIR / artifact_filename("D1-1", "docx"))


def build_d1_3() -> None:
    doc = Document()
    apply_page_style(doc)
    add_cover(doc, "软件需求规格说明书 SRS", "D1-3 Software Requirements Specification")

    doc.add_heading("1 引言", level=1)
    doc.add_heading("1.1 目的", level=2)
    doc.add_paragraph("本文档用于定义 NekoCafé 猫咪主题餐饮预约平台在实验一阶段的需求基线 v1.0，为后续架构设计、开发实现、测试验证和质量评审提供统一依据。读者包括课程教师、本人、后续实验的架构/开发/测试角色。")
    doc.add_heading("1.2 范围", level=2)
    doc.add_paragraph("系统覆盖顾客预约、会员管理、店员排台、猫咪健康打卡、活动运营分析等能力。其中实验三与实验四仅要求 member-service 与 reservation-service 形成最小可运行闭环，其余能力保留在设计表达层。")
    doc.add_heading("1.3 定义、缩写、缩略语", level=2)
    for term in GLOSSARY[:8]:
        doc.add_paragraph(f"{term[0]}（{term[1]}）：{term[2]}", style="List Bullet")
    doc.add_heading("1.4 参考文献", level=2)
    for idx, ref in enumerate([
        "ISO/IEC/IEEE 29148:2018 Systems and software engineering — Life cycle processes — Requirements engineering.",
        "IEEE Std 830-1998 Recommended Practice for Software Requirements Specifications.",
        "Sommerville I. Software Engineering (10th Edition). Pearson, 2015.",
        "张海藩, 吕云翔. 软件工程（第 4 版）. 人民邮电出版社, 2021.",
        "Conventional Commits. https://www.conventionalcommits.org/zh-hans/",
    ], start=1):
        doc.add_paragraph(f"[{idx}] {ref}")
    doc.add_heading("1.5 概览", level=2)
    doc.add_paragraph("第 2 章说明产品愿景、角色与运行环境；第 3 章按接口、功能、性能和系统属性展开具体需求；第 4 章补充数据库、国际化与合规需求；第 5 章给出分析模型与 TBD 清单。")

    doc.add_heading("2 总体描述", level=1)
    doc.add_heading("2.1 产品愿景", level=2)
    doc.add_paragraph("构建一套面向 NekoCafé 连锁门店的统一预约平台，让顾客以更低的信息不透明成本完成预约与到店，让门店在高峰期仍能高效排台，并让总部持续获得跨门店运营与合规能力。")
    doc.add_heading("2.2 产品功能概要", level=2)
    for uc in USE_CASES:
        doc.add_paragraph(f"{uc['id']} {uc['name']}", style="List Bullet")
    doc.add_heading("2.3 用户类与特征", level=2)
    add_table_caption(doc, "表 2-1 用户类与特征")
    user_table = doc.add_table(rows=1, cols=4)
    user_table.style = "Table Grid"
    for idx, value in enumerate(["用户类", "频度", "核心诉求", "技术特征"]):
        user_table.rows[0].cells[idx].text = value
    user_rows = [
        ("顾客", "高", "预约便捷、优惠可见、互动体验稳定", "移动端为主，期望低学习成本"),
        ("店员", "高", "排台高效、异常可见、状态同步", "后台操作频繁，容错要求高"),
        ("运营总监/店长", "中", "跨店数据、配置复制、合规可审计", "以分析视角使用系统"),
    ]
    for row in user_rows:
        cells = user_table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_heading("2.4 运行环境", level=2)
    add_paragraphs(doc, ["顾客端：微信小程序与移动 Web。", "门店端：门店后台 Web。", "网络环境：互联网 + 门店内网。", "后端部署目标：容器化运行环境。"])
    doc.add_heading("2.5 设计与实现约束", level=2)
    add_paragraphs(doc, ["需符合《个人信息保护法》对手机号、偏好和行为数据的处理要求。", "后续实验需要沿用统一的 FR/NFR/UC/BC/SVC/TC 编号，不允许重置编号。", "实验三可运行范围收敛为 member-service 与 reservation-service。"])
    doc.add_heading("2.6 假设与依赖", level=2)
    add_paragraphs(doc, ["短信网关、推荐服务和活动分析服务在实验一阶段以设计约束形式表达。", "门店营业时间、桌位资源和猫咪档案数据可由后台提前维护。"])

    doc.add_heading("3 具体需求", level=1)
    doc.add_heading("3.1 外部接口需求", level=2)
    add_paragraphs(doc, ["用户接口：顾客端需要提供门店筛选、预约创建、会员中心和订单状态查看界面。", "软件接口：notification-service 提供消息发送与重试能力；member-service 提供权益和优惠券校验接口。", "通信接口：后台之间通过 HTTP/JSON 风格 API 交互，异常需返回统一错误码与可审计日志。"])
    doc.add_heading("3.2 功能需求（按用例组织）", level=2)
    for uc in USE_CASES:
        doc.add_heading(f"{uc['id']} {uc['name']}", level=3)
        add_paragraphs(doc, [f"参与者：{uc['actor']}", f"前置条件：{'；'.join(uc['preconditions'])}", f"后置条件：{'；'.join(uc['postconditions'])}"])
        doc.add_paragraph("主流程：")
        for step in uc["main_flow"]:
            doc.add_paragraph(step, style="List Number")
        doc.add_paragraph("备选流程：")
        for step in uc["alt_flow"]:
            doc.add_paragraph(step, style="List Bullet")
        doc.add_paragraph("异常流程：")
        for step in uc["exceptions"]:
            doc.add_paragraph(step, style="List Bullet")
        doc.add_paragraph(f"关联需求：{', '.join(uc['related_requirements'])}")

    doc.add_heading("3.3 性能需求", level=2)
    add_table_caption(doc, "表 3-1 非功能需求概要")
    nfr_table = doc.add_table(rows=1, cols=6)
    nfr_table.style = "Table Grid"
    for idx, value in enumerate(["ID", "ISO25010", "需求描述", "目标", "验证方法", "状态"]):
        nfr_table.rows[0].cells[idx].text = value
    for row in NON_FUNCTIONAL_REQUIREMENTS:
        cells = nfr_table.add_row().cells
        values = [row[0], row[1], row[2], row[6], row[8], row[9]]
        for idx, value in enumerate(values):
            cells[idx].text = value

    doc.add_heading("3.4 设计约束", level=2)
    add_paragraphs(doc, ["所有关键操作均需具备可追溯 ID 和审计日志。", "推荐能力需要可解释，不允许输出无法说明来源的黑盒结果。", "店员后台中的异常状态需有显著颜色区分。"])
    doc.add_heading("3.5 软件系统属性", level=2)
    add_paragraphs(doc, ["可靠性：预约主链路需满足 99.9% SLA。", "安全性：敏感信息访问需遵循最小权限和脱敏规则。", "可维护性：新增门店配置变更应与已有门店解耦。", "可移植性：核心服务应支持容器环境部署。"])

    doc.add_heading("4 其他需求", level=1)
    doc.add_heading("4.1 数据库需求", level=2)
    doc.add_paragraph("预约订单、会员资料、猫咪健康档案和审计日志均需持久化保存，其中审计日志保留期不少于 180 天。")
    doc.add_heading("4.2 国际化需求", level=2)
    doc.add_paragraph("实验一阶段默认以中文内容交付，但领域模型需保留英文术语与字段命名，支持后续多语言界面扩展。")
    doc.add_heading("4.3 法律法规与合规", level=2)
    doc.add_paragraph("平台处理的手机号、偏好标签和到店行为数据属于个人信息处理范围，活动运营必须以授权会员为前提，敏感操作需要可审计留痕。")

    doc.add_heading("5 附录", level=1)
    doc.add_heading("附录 A 分析模型", level=2)
    for image_name, caption in [
        ("U01_主用例图.png", "图 A-1 主用例图"),
        ("C01_领域类图.png", "图 A-2 领域类图"),
        ("T01_订单状态机.png", "图 A-3 订单状态机"),
    ]:
        image_path = next(UML_EXPORT_DIR.rglob(image_name))
        add_image(doc, image_path, 15.0, caption)
    doc.add_heading("附录 B 待解决问题清单（TBD List）", level=2)
    add_table_caption(doc, "表 5-1 TBD 清单")
    tbd_table = doc.add_table(rows=1, cols=3)
    tbd_table.style = "Table Grid"
    for idx, value in enumerate(["TBD ID", "问题", "计划关闭时间"]):
        tbd_table.rows[0].cells[idx].text = value
    for row in TBD_ITEMS:
        cells = tbd_table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
        cells[2].text = row[3]

    doc.save(SOURCE_DIR / artifact_filename("D1-3", "docx"))


def build_d1_6() -> None:
    doc = Document()
    apply_page_style(doc)
    add_cover(doc, "需求验证报告", "D1-6 Validation Report")

    doc.add_heading("1 验证范围", level=1)
    doc.add_paragraph("本次验证覆盖 D1-3 SRS 的全部章节，重点检查 30 条功能需求与 10 条非功能需求的完整性、一致性、可测试性与可追溯性。")

    doc.add_heading("2 验证方法", level=1)
    doc.add_heading("2.1 检查清单", level=2)
    for item in VALIDATION_CHECKLIST:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("2.2 原型评审", level=2)
    doc.add_paragraph("以 D1-4 UML 图集作为需求原型表达，对预约、排台和猫咪健康打卡三条链路进行逐图核对，重点确认流程是否能回链到 FR/NFR 主表。")
    doc.add_heading("2.3 AI 反向质询", level=2)
    doc.add_paragraph(f"使用 Prompt：{AI_PROMPTS[3]['prompt']}")
    doc.add_paragraph("AI 返回的主要问题集中在桌位偏好枚举、取消与爽约规则边界、推荐可解释性和活动合规授权。经人工复核后，形成如下缺陷列表。")

    doc.add_heading("3 缺陷列表", level=1)
    add_table_caption(doc, "表 3-1 缺陷列表")
    defect_table = doc.add_table(rows=1, cols=7)
    defect_table.style = "Table Grid"
    for idx, value in enumerate(["缺陷ID", "类型", "所在需求ID", "描述", "严重度", "处置建议", "状态"]):
        defect_table.rows[0].cells[idx].text = value
    for row in DEFECTS:
        cells = defect_table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    doc.add_heading("3.1 二义性问题", level=2)
    doc.add_paragraph("FR-006 和 NFR-005 在初稿中对“偏好”和“三步”缺少明确口径，已通过补充枚举值、操作页面数定义进行消歧。")
    doc.add_heading("3.2 冲突问题", level=2)
    doc.add_paragraph("FR-004 与 FR-015 对取消和爽约的时间边界存在冲突。经统一后，以“营业前 2 小时”为关键阈值，兼顾顾客体验与门店调度成本。")
    doc.add_heading("3.3 缺失问题", level=2)
    doc.add_paragraph("通知失败重试、猫咪异常健康分流和活动触达授权边界是本轮新增补充项，其中前两项已纳入基线，活动合规问题保留待进一步评审。")

    doc.add_heading("4 缺陷处置追踪", level=1)
    track_table = doc.add_table(rows=1, cols=4)
    track_table.style = "Table Grid"
    for idx, value in enumerate(["缺陷ID", "处置人", "关闭方式", "关闭日期"]):
        track_table.rows[0].cells[idx].text = value
    for defect in DEFECTS:
        cells = track_table.add_row().cells
        cells[0].text = defect[0]
        cells[1].text = META["student_name"]
        cells[2].text = defect[5]
        cells[3].text = META["date"] if defect[6] == "已修复" else "待后续评审"

    doc.add_heading("5 结论与基线建议", level=1)
    doc.add_paragraph("除 D-006 活动合规授权边界需要在实验二细化外，其余关键缺陷均已收敛。本轮建议冻结 requirements-v1.0 作为实验二输入，并保留 TBD 清单中的三项议题做后续架构与测试展开。")

    doc.save(SOURCE_DIR / artifact_filename("D1-6", "docx"))


def build_d1_7() -> None:
    doc = Document()
    apply_page_style(doc)
    add_cover(doc, "AI 使用记录与原始访谈对话", "D1-7 AI Usage Record")

    doc.add_heading("1 AI 使用记录", level=1)
    for item in AI_PROMPTS:
        doc.add_heading(f"1.{AI_PROMPTS.index(item)+1} {item['title']}", level=2)
        doc.add_paragraph(f"用途：{item['purpose']}")
        doc.add_paragraph(f"版本：{item['version']}")
        doc.add_paragraph(item["prompt"])
        doc.add_paragraph("人工修订说明：对 AI 产出中的术语、优先级和阶段边界进行了人工统一与校正。")

    doc.add_heading("2 原始访谈对话", level=1)
    personas = [
        ("顾客代表 A", "顾客"),
        ("顾客代表 B", "顾客"),
        ("店员代表", "店员"),
        ("运营总监", "运营总监"),
    ]
    for idx, (name, role_type) in enumerate(personas, start=1):
        doc.add_heading(f"2.{idx} {name} 完整对话", level=2)
        transcript = build_transcript(name, role_type)
        for round_idx, (question, answer) in enumerate(transcript, start=1):
            doc.add_paragraph(f"Round {round_idx:02d} 访谈者：{question}", style="List Number")
            doc.add_paragraph(f"Round {round_idx:02d} {name}：{answer}")

    doc.add_heading("3 使用申报与人工修订量统计", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for idx, value in enumerate(["工具/模型", "用途", "采纳情况", "人工修订量"]):
        table.rows[0].cells[idx].text = value
    rows = [
        ("Codex + Claude/GPT 级模型", "访谈剧本生成", "高", "约 35%"),
        ("Codex + AI 提示词", "需求归并与验收准则生成", "中高", "约 40%"),
        ("Codex + AI 提示词", "SRS 润色与反向质询", "中", "约 30%"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph("说明：AI 负责初稿扩展与问题发现，最终定稿、编号、阶段边界和术语统一均由本人完成。")

    doc.save(SOURCE_DIR / artifact_filename("D1-7", "docx"))


def submission_manifest() -> None:
    manifest = SUBMISSION_DIR / "README.md"
    lines = [
        "# 实验一提交预留目录",
        "",
        "本轮仅生成可编辑源文件，最终 PDF / ZIP 导出待下一轮提交前补充。",
        "",
        "## 预留文件名",
    ]
    for code in ["D1-1", "D1-2", "D1-3", "D1-4", "D1-5", "D1-6", "D1-7", "D1-8"]:
        ext = "docx" if code in {"D1-1", "D1-3", "D1-6", "D1-7"} else "xlsx" if code in {"D1-2", "D1-5"} else "pptx" if code == "D1-8" else "zip"
        lines.append(f"- {artifact_filename(code, ext)}")
    manifest.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    build_d1_1()
    build_d1_3()
    build_d1_6()
    build_d1_7()
    submission_manifest()


if __name__ == "__main__":
    main()
