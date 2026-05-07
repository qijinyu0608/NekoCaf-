from __future__ import annotations

import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document

from build_lab2_docs import add_cover, add_image, add_table_caption, apply_page_style
from lab2_batch2_common import (
    C4_VIEWS,
    CROSS_BORDER_RULES,
    DATA_LIFECYCLE,
    DATA_MODEL_CONTEXTS,
    INDEX_AND_PARTITION,
    MULTI_TENANT_STRATEGY,
    OPENAPI_SPECS,
    d2_4_export_dir,
    d2_5_package_dir,
    d2_6_export_dir,
    d2_6_filename,
)
from lab2_common import DIAGRAM_DIR, META, REFERENCES, SOURCE_DIR, artifact_stem, ensure_dirs


FONT_PATH = "/System/Library/Fonts/PingFang.ttc"
C4_DIAGRAM_DIR = DIAGRAM_DIR / artifact_stem("D2-4")
ER_EXPORT_DIR = d2_6_export_dir()
OPENAPI_DIR = d2_5_package_dir()


def font(size: int):
    try:
        return ImageFont.truetype(FONT_PATH, size)
    except Exception:
        return ImageFont.load_default()


def write_json_yaml(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_contract_preview(path: Path, title: str, description: str, endpoints: list[str]) -> None:
    img = Image.new("RGB", (1440, 900), "#F7F8FA")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((60, 60, 1380, 840), radius=24, fill="#FFFFFF", outline="#CBD5E1", width=3)
    draw.rectangle((60, 60, 1380, 130), fill="#1F2937")
    draw.text((95, 86), f"Swagger Contract Preview - {title}", fill="#FFFFFF", font=font(28))
    draw.text((95, 170), description, fill="#111827", font=font(22))
    y = 260
    for endpoint in endpoints:
        draw.rounded_rectangle((95, y, 1310, y + 82), radius=16, fill="#EEF2FF", outline="#C7D2FE", width=2)
        draw.text((120, y + 24), endpoint, fill="#1E3A8A", font=font(24))
        y += 110
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def normalize_png_for_docx(path: Path) -> Path:
    normalized = path.with_name(f"{path.stem}_docx.png")
    Image.open(path).save(normalized, format="PNG")
    return normalized


def build_d2_4_readme() -> None:
    doc = Document()
    apply_page_style(doc)
    add_cover(doc, "C4 架构图集 README", "D2-4 C4 Model Set README", "D2-4")

    doc.add_heading("1 本目录用途", level=1)
    doc.add_paragraph("本目录存放实验二 C4 四层视图的 drawio 源文件与 PNG 导出图。图集围绕“微服务主线 + 核心双服务先落地”的结论组织，便于在答辩时逐层讲解。")

    doc.add_heading("2 图集清单", level=1)
    add_table_caption(doc, "表 2-1 C4 视图清单")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["视图编号", "文件名", "层级", "用途"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for view in C4_VIEWS:
        row = table.add_row().cells
        row[0].text = view["id"]
        row[1].text = view["file"]
        row[2].text = view["title"]
        row[3].text = view["goal"]

    doc.add_heading("3 图形约定", level=1)
    bullets = [
        "人物与外部系统使用浅蓝色或浅绿色块表示，核心服务边界使用暖色块突出。",
        "容器层将实验三会优先落地的 member-service 与 reservation-service 放在中央强调位置。",
        "所有图均导出为 PNG，并在图集目录内同步保留 exports 子目录，便于按模板核查。",
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("4 与实验一/实验三的衔接", level=1)
    doc.add_paragraph("L1-L4 视图与 D2-1 的高优先级 QAS、D2-2 的六个限界上下文、D2-3 的 ATAM 结论一一对应；同时 L3/L4 明确聚焦 member-service 和 reservation-service，为实验三服务代码实现保留直接输入。")

    doc.save(C4_DIAGRAM_DIR / "README.docx")


def build_d2_5_package() -> None:
    OPENAPI_DIR.mkdir(parents=True, exist_ok=True)
    (OPENAPI_DIR / "common").mkdir(parents=True, exist_ok=True)
    (OPENAPI_DIR / "examples").mkdir(parents=True, exist_ok=True)
    (OPENAPI_DIR / "screenshots").mkdir(parents=True, exist_ok=True)

    for service_name, spec in OPENAPI_SPECS.items():
        write_json_yaml(OPENAPI_DIR / f"{service_name}.yaml", spec)

    common_errors = {
        "MEMBER_NOT_FOUND": "会员不存在或不属于当前租户",
        "RESERVATION_NOT_FOUND": "预约不存在或已被归档",
        "SLOT_CONFLICT": "桌位资源已被其他请求抢占",
        "CAT_RESTRICTION": "猫咪互动限制阻断当前预约",
        "FORBIDDEN": "鉴权通过但无资源访问权限",
    }
    write_json_yaml(OPENAPI_DIR / "common" / "error-codes.json", common_errors)

    write_json_yaml(
        OPENAPI_DIR / "examples" / "member-profile-update.json",
        {"nickname": "Momo", "preferences": ["安静角落", "靠窗"], "allergyNote": "不接触猫薄荷"},
    )
    write_json_yaml(
        OPENAPI_DIR / "examples" / "create-reservation.json",
        {
            "memberId": "MBR-1001",
            "storeId": "STORE-BJFY-01",
            "slotId": "SLOT-20260508-1900-A08",
            "partySize": 2,
            "preferredTheme": "治愈互动区",
            "catInteractionMode": "LOW_STIMULATION",
        },
    )

    write_contract_preview(
        OPENAPI_DIR / "screenshots" / "member-service-preview.png",
        "member-service",
        "会员服务契约概览，覆盖资料、偏好、积分与优惠券接口。",
        ["GET /member/v1/members/{memberId}", "PATCH /member/v1/members/{memberId}/profile", "GET /member/v1/members/{memberId}/coupons"],
    )
    write_contract_preview(
        OPENAPI_DIR / "screenshots" / "reservation-service-preview.png",
        "reservation-service",
        "预约服务契约概览，覆盖时段查询、创建预约、查询详情与取消。",
        ["GET /reservation/v1/stores/{storeId}/slots", "POST /reservation/v1/reservations", "GET /reservation/v1/reservations/{reservationId}", "POST /reservation/v1/reservations/{reservationId}/cancel"],
    )

    doc = Document()
    apply_page_style(doc)
    add_cover(doc, "OpenAPI 契约文档 README", "D2-5 OpenAPI Contracts README", "D2-5")

    doc.add_heading("1 本目录用途", level=1)
    doc.add_paragraph("本目录存放桌位预约服务与会员服务的 OpenAPI 3.0 契约、通用错误码、示例载荷和契约概览图。考虑到当前环境缺少 swagger-cli，本批次采用 JSON 兼容 YAML 形式写入 .yaml 文件，并通过结构化检查保证关键段落完整。")

    doc.add_heading("2 文件结构", level=1)
    add_table_caption(doc, "表 2-1 D2-5 文件结构")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, header in enumerate(["路径", "内容", "备注"]):
        table.rows[0].cells[idx].text = header
    rows = [
        ("member-service.yaml", "会员服务主契约", "包含 info.description、servers、securitySchemes、responses"),
        ("reservation-service.yaml", "预约服务主契约", "包含限流扩展、错误码与双核心服务接口"),
        ("common/error-codes.json", "通用错误码字典", "便于实验三和实验四复用"),
        ("examples/*.json", "请求示例", "可直接作为联调输入"),
        ("screenshots/*.png", "契约概览图", "本批次用作 Swagger UI 预览占位"),
    ]
    for path, content, remark in rows:
        row = table.add_row().cells
        row[0].text = path
        row[1].text = content
        row[2].text = remark

    doc.add_heading("3 契约红线", level=1)
    bullets = [
        "所有接口均带版本化 path 前缀和 servers 根地址，便于后续灰度与多环境发布。",
        "所有核心操作均声明 bearerAuth 与 X-Tenant-Id，用于表达鉴权和租户隔离。",
        "所有操作至少带一个 x-ratelimit-* 扩展，保证限流策略可审阅。",
        "所有错误响应统一回落到 components.responses，避免实验三接口各自发散。",
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.save(OPENAPI_DIR / "README.docx")


def build_d2_6_doc() -> None:
    doc = Document()
    apply_page_style(doc)
    add_cover(doc, "数据模型 ER 图与数据字典", "D2-6 ER Diagram & Data Dictionary", "D2-6")

    core_png = ER_EXPORT_DIR / "ER01_CoreReservation.png"
    store_png = ER_EXPORT_DIR / "ER02_StoreCatOps.png"
    reco_png = ER_EXPORT_DIR / "ER03_RecommendationOps.png"
    core_docx_png = normalize_png_for_docx(core_png)
    store_docx_png = normalize_png_for_docx(store_png)
    reco_docx_png = normalize_png_for_docx(reco_png)

    doc.add_heading("1 ER 图（按上下文分图）", level=1)
    add_image(doc, core_docx_png, 15.2, "图 1-1 核心会员-预约-订单 ER 图")
    add_image(doc, store_docx_png, 15.2, "图 1-2 门店运营与猫咪健康 ER 图")
    add_image(doc, reco_docx_png, 15.2, "图 1-3 推荐与运营支撑 ER 图")

    doc.add_heading("2 数据字典", level=1)
    for context in DATA_MODEL_CONTEXTS:
        doc.add_paragraph(f"{context['context']}（{context['tenant_strategy']}）", style="Heading 2")
        for entity in context["entities"]:
            add_table_caption(doc, f"表 {entity['name']} 字段字典")
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for idx, header in enumerate(["字段名", "类型", "含义", "所属实体"]):
                table.rows[0].cells[idx].text = header
            for field_name, field_type, meaning in entity["fields"]:
                row = table.add_row().cells
                row[0].text = field_name
                row[1].text = field_type
                row[2].text = meaning
                row[3].text = entity["name"]

    doc.add_heading("3 索引与分区策略", level=1)
    for item in INDEX_AND_PARTITION:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4 数据生命周期", level=1)
    add_table_caption(doc, "表 4-1 数据生命周期")
    lifecycle = doc.add_table(rows=1, cols=3)
    lifecycle.style = "Table Grid"
    for idx, header in enumerate(["实体", "保留策略", "说明"]):
        lifecycle.rows[0].cells[idx].text = header
    for entity, duration, note in DATA_LIFECYCLE:
        row = lifecycle.add_row().cells
        row[0].text = entity
        row[1].text = duration
        row[2].text = note

    doc.add_heading("5 多租户隔离策略", level=1)
    doc.add_paragraph(MULTI_TENANT_STRATEGY)

    doc.add_heading("6 跨境数据合规标注", level=1)
    doc.add_paragraph("以下规则用于解释所有 [GDPR] 字段的跨境治理要求：")
    for item in CROSS_BORDER_RULES:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7 参考文献", level=1)
    for item in REFERENCES:
        doc.add_paragraph(item)

    doc.save(SOURCE_DIR / d2_6_filename())


if __name__ == "__main__":
    ensure_dirs()
    C4_DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    OPENAPI_DIR.mkdir(parents=True, exist_ok=True)
    build_d2_4_readme()
    build_d2_5_package()
    build_d2_6_doc()
