from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from docx_style_policy import apply_black_text_policy
from lab2_common import (
    ARCHITECTURE_DRIVERS,
    ATAM_ANALYSIS,
    BOUNDED_CONTEXTS,
    CANDIDATE_ARCHITECTURES,
    COMMAND_AGGREGATE_VIEW,
    CONTEXT_RELATIONS,
    DOMAIN_EVENTS,
    EXPORT_DIR,
    META,
    QUALITY_ATTRIBUTE_SCENARIOS,
    REFERENCES,
    SOURCE_DIR,
    UTILITY_TREE,
    artifact_filename,
    artifact_stem,
    ensure_dirs,
    requirement_brief,
)


LAB2_EXPORT_DIR = EXPORT_DIR / artifact_stem("D2-2")


def apply_page_style(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section)

    normal = doc.styles["Normal"]
    apply_black_text_policy(doc)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def configure_section(section) -> None:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.text = META["header"]

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)


def add_page_number(paragraph) -> None:
    paragraph.text = ""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def add_cover(doc: Document, title: str, subtitle: str, artifact_code: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    p.style = doc.styles["Title"]

    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"
    pairs = [
        ("课程 / 实验", f"{META['course']} / {META['experiment']} {META['experiment_title']}"),
        ("产出编号", artifact_code),
        ("班级 / 学号 / 姓名", f"{META['class_name']} / {META['student_id']} / {META['student_name']}"),
        ("版本 / 日期", f"{META['version']} / {META['date']}"),
    ]
    for row_idx, pair in enumerate(pairs):
        info.rows[row_idx].cells[0].text = pair[0]
        info.rows[row_idx].cells[1].text = pair[1]
        for cell in info.rows[row_idx].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("学术诚信声明", style="Heading 2")
    doc.add_paragraph(META["declaration"])
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(new_section)


def add_table_caption(doc: Document, caption: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.add_run(caption).bold = True


def add_image(doc: Document, image_path: Path, width_cm: float, caption: str) -> None:
    doc.add_picture(str(image_path), width=Cm(width_cm))
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run(caption).bold = True


def build_d2_1() -> None:
    doc = Document()
    apply_page_style(doc)
    add_cover(doc, "质量属性场景与效用树", "D2-1 Quality Attribute Scenarios & Utility Tree", "D2-1")

    doc.add_heading("1 项目背景与架构驱动力", level=1)
    doc.add_paragraph("本报告承接实验一需求基线，从 10 条非功能需求中提炼实验二的架构驱动力，重点回答高峰期 5000 QPS、零停机扩店与会员数据合规三类核心问题。")
    for item in ARCHITECTURE_DRIVERS:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2 效用树", level=1)
    add_table_caption(doc, "表 2-1 效用树根节点与优先级")
    utility_table = doc.add_table(rows=1, cols=4)
    utility_table.style = "Table Grid"
    for idx, value in enumerate(["根节点", "二级维度", "叶子节点(QAS)", "优先级"]):
        utility_table.rows[0].cells[idx].text = value
    for root_name, dimension, leaves, priority in UTILITY_TREE:
        cells = utility_table.add_row().cells
        cells[0].text = root_name
        cells[1].text = dimension
        cells[2].text = leaves
        cells[3].text = priority

    doc.add_heading("3 Top-10 QAS 详表", level=1)
    add_table_caption(doc, "表 3-1 Top-10 质量属性场景")
    qas_table = doc.add_table(rows=1, cols=9)
    qas_table.style = "Table Grid"
    headers = ["场景ID", "来源NFR", "刺激源", "刺激", "制品", "环境", "响应", "响应度量", "优先级"]
    for idx, value in enumerate(headers):
        qas_table.rows[0].cells[idx].text = value
    for qas in QUALITY_ATTRIBUTE_SCENARIOS:
        cells = qas_table.add_row().cells
        cells[0].text = qas["id"]
        cells[1].text = qas["source_nfr"]
        cells[2].text = qas["stimulus_source"]
        cells[3].text = qas["stimulus"]
        cells[4].text = qas["artifact"]
        cells[5].text = qas["environment"]
        cells[6].text = qas["response"]
        cells[7].text = qas["response_measure"]
        cells[8].text = qas["priority"]

    doc.add_heading("4 架构关注点清单", level=1)
    focus_points = [
        "预约主链路必须和通知、推荐等扩展能力解耦，避免非核心能力拖慢高峰期性能。",
        "门店复制能力应优先通过配置、模板和部署标准化解决，而不是人为运维操作。",
        "会员与运营敏感数据需要独立治理，保证权限控制、脱敏和审计链条清晰。",
        "实验三只落地双核心服务，因此实验二结论必须兼顾“完整表达”与“可控实现”。",
    ]
    for item in focus_points:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5 优先级裁决结论", level=1)
    doc.add_paragraph("本轮裁决结果将性能效率、可靠性、合规安全和零停机扩店列为最高优先级；顾客体验与推荐解释性作为第二层驱动。基于这一结论，实验二后续架构方案优先偏向微服务边界表达，但实验三只承诺先落地 member-service 与 reservation-service。")
    doc.add_heading("6 参考文献", level=1)
    for item in REFERENCES:
        doc.add_paragraph(item)

    doc.save(SOURCE_DIR / artifact_filename("D2-1", "docx"))


def build_d2_2() -> None:
    doc = Document()
    apply_page_style(doc)
    add_cover(doc, "事件风暴与上下文映射图", "D2-2 Event Storming & Context Map", "D2-2")

    event_png = LAB2_EXPORT_DIR / "ES01_事件风暴图.png"
    context_png = LAB2_EXPORT_DIR / "CM01_上下文映射图.png"

    doc.add_heading("1 方法说明", level=1)
    doc.add_paragraph("本报告采用事件风暴方法从实验一 FR/NFR 与 UML 产物中抽取领域事件，再以 DDD 限界上下文组织模型边界。输出既服务于实验二架构设计，也为实验三核心服务落地保留清晰边界。")

    doc.add_heading("2 领域事件列表", level=1)
    add_table_caption(doc, "表 2-1 领域事件清单")
    events_table = doc.add_table(rows=1, cols=3)
    events_table.style = "Table Grid"
    for idx, value in enumerate(["序号", "领域事件", "归属业务域"]):
        events_table.rows[0].cells[idx].text = value
    for index, (event_name, domain) in enumerate(DOMAIN_EVENTS, start=1):
        cells = events_table.add_row().cells
        cells[0].text = f"E-{index:02d}"
        cells[1].text = event_name
        cells[2].text = domain

    doc.add_heading("3 命令、聚合、策略与读模型", level=1)
    add_table_caption(doc, "表 3-1 聚合识别结果")
    cap_table = doc.add_table(rows=1, cols=4)
    cap_table.style = "Table Grid"
    for idx, value in enumerate(["命令", "聚合", "策略/约束", "读模型"]):
        cap_table.rows[0].cells[idx].text = value
    for row in COMMAND_AGGREGATE_VIEW:
        cells = cap_table.add_row().cells
        cells[0].text = row["command"]
        cells[1].text = row["aggregate"]
        cells[2].text = row["policy"]
        cells[3].text = row["read_model"]

    doc.add_heading("4 限界上下文清单", level=1)
    add_table_caption(doc, "表 4-1 限界上下文与追溯关系")
    bc_table = doc.add_table(rows=1, cols=6)
    bc_table.style = "Table Grid"
    for idx, value in enumerate(["BC", "名称", "服务归属", "目标", "关联需求", "关键聚合"]):
        bc_table.rows[0].cells[idx].text = value
    for context in BOUNDED_CONTEXTS:
        cells = bc_table.add_row().cells
        cells[0].text = context["id"]
        cells[1].text = context["name"]
        cells[2].text = context["service"]
        cells[3].text = context["goal"]
        cells[4].text = "\n".join(f"{req}: {requirement_brief(req)}" for req in context["requirements"])
        cells[5].text = "\n".join(context["aggregates"])

    doc.add_heading("5 关键聚合不变式", level=1)
    for context in BOUNDED_CONTEXTS:
        doc.add_paragraph(f"{context['id']} {context['name']}", style="Heading 3")
        for invariant in context["invariants"]:
            doc.add_paragraph(invariant, style="List Bullet")

    doc.add_heading("6 事件风暴图", level=1)
    add_image(doc, event_png, 15.0, "图 6-1 事件风暴图")
    doc.add_paragraph("事件风暴图以业务时间线展开，从顾客注册、预约创建到门店运营、猫咪健康和推荐结果生成，突出核心业务事件与跨域协作路径。")

    doc.add_heading("7 上下文映射图", level=1)
    add_image(doc, context_png, 15.0, "图 7-1 上下文映射图")
    add_table_caption(doc, "表 7-1 上下文关系说明")
    relation_table = doc.add_table(rows=1, cols=4)
    relation_table.style = "Table Grid"
    for idx, value in enumerate(["上游上下文", "下游上下文", "关系类型", "说明"]):
        relation_table.rows[0].cells[idx].text = value
    for source, target, relation, note in CONTEXT_RELATIONS:
        cells = relation_table.add_row().cells
        cells[0].text = source
        cells[1].text = target
        cells[2].text = relation
        cells[3].text = note

    doc.add_heading("8 结论", level=1)
    doc.add_paragraph("本轮将会员、预约、订单、门店运营、猫咪健康、推荐六个上下文作为实验二的主边界，其中 member-service 与 reservation-service 作为实验三优先落地的核心服务，其余上下文保留在架构与契约设计层。")
    doc.add_heading("9 参考文献", level=1)
    for item in REFERENCES:
        doc.add_paragraph(item)

    doc.save(SOURCE_DIR / artifact_filename("D2-2", "docx"))


def build_d2_3() -> None:
    doc = Document()
    apply_page_style(doc)
    add_cover(doc, "候选架构对比与 ATAM 报告", "D2-3 Architecture Alternatives & ATAM", "D2-3")

    doc.add_heading("1 商业目标与评审背景", level=1)
    doc.add_paragraph("实验二需要回答三类核心问题：高峰期 5000 QPS 怎么扛、新门店如何零停机接入、会员数据跨境同步如何合规。因此本报告基于实验一 QAS 与 DDD 边界，对两套候选架构进行 ATAM 风格对比。")
    for goal in ATAM_ANALYSIS["business_goals"]:
        doc.add_paragraph(goal, style="List Bullet")

    doc.add_heading("2 候选方案 A：模块化单体", level=1)
    doc.add_paragraph(CANDIDATE_ARCHITECTURES[0]["summary"])
    add_table_caption(doc, "表 2-1 方案 A 优势与代价")
    plan_a = doc.add_table(rows=1, cols=2)
    plan_a.style = "Table Grid"
    plan_a.rows[0].cells[0].text = "优势"
    plan_a.rows[0].cells[1].text = "代价"
    max_rows = max(len(CANDIDATE_ARCHITECTURES[0]["strengths"]), len(CANDIDATE_ARCHITECTURES[0]["tradeoffs"]))
    for idx in range(max_rows):
        cells = plan_a.add_row().cells
        cells[0].text = CANDIDATE_ARCHITECTURES[0]["strengths"][idx] if idx < len(CANDIDATE_ARCHITECTURES[0]["strengths"]) else ""
        cells[1].text = CANDIDATE_ARCHITECTURES[0]["tradeoffs"][idx] if idx < len(CANDIDATE_ARCHITECTURES[0]["tradeoffs"]) else ""

    doc.add_heading("3 候选方案 B：微服务 + 事件驱动", level=1)
    doc.add_paragraph(CANDIDATE_ARCHITECTURES[1]["summary"])
    add_table_caption(doc, "表 3-1 方案 B 优势与代价")
    plan_b = doc.add_table(rows=1, cols=2)
    plan_b.style = "Table Grid"
    plan_b.rows[0].cells[0].text = "优势"
    plan_b.rows[0].cells[1].text = "代价"
    max_rows = max(len(CANDIDATE_ARCHITECTURES[1]["strengths"]), len(CANDIDATE_ARCHITECTURES[1]["tradeoffs"]))
    for idx in range(max_rows):
        cells = plan_b.add_row().cells
        cells[0].text = CANDIDATE_ARCHITECTURES[1]["strengths"][idx] if idx < len(CANDIDATE_ARCHITECTURES[1]["strengths"]) else ""
        cells[1].text = CANDIDATE_ARCHITECTURES[1]["tradeoffs"][idx] if idx < len(CANDIDATE_ARCHITECTURES[1]["tradeoffs"]) else ""

    doc.add_heading("4 ATAM 评估", level=1)
    add_table_caption(doc, "表 4-1 QAS 与方案匹配")
    qas_table = doc.add_table(rows=1, cols=4)
    qas_table.style = "Table Grid"
    for idx, value in enumerate(["QAS", "质量属性", "方案 A", "方案 B"]):
        qas_table.rows[0].cells[idx].text = value
    plan_a_qas = set(CANDIDATE_ARCHITECTURES[0]["qas_fit"])
    plan_b_qas = set(CANDIDATE_ARCHITECTURES[1]["qas_fit"])
    for qas in QUALITY_ATTRIBUTE_SCENARIOS:
        cells = qas_table.add_row().cells
        cells[0].text = qas["id"]
        cells[1].text = qas["quality"]
        cells[2].text = "匹配" if qas["id"] in plan_a_qas else "弱"
        cells[3].text = "匹配" if qas["id"] in plan_b_qas else "弱"

    doc.add_heading("4.1 敏感点", level=2)
    for item in ATAM_ANALYSIS["sensitivity_points"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("4.2 权衡点", level=2)
    for item in ATAM_ANALYSIS["tradeoff_points"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("4.3 风险与非风险", level=2)
    add_table_caption(doc, "表 4-2 风险与非风险")
    risk_table = doc.add_table(rows=1, cols=2)
    risk_table.style = "Table Grid"
    risk_table.rows[0].cells[0].text = "风险"
    risk_table.rows[0].cells[1].text = "非风险"
    max_rows = max(len(ATAM_ANALYSIS["risks"]), len(ATAM_ANALYSIS["non_risks"]))
    for idx in range(max_rows):
        cells = risk_table.add_row().cells
        cells[0].text = ATAM_ANALYSIS["risks"][idx] if idx < len(ATAM_ANALYSIS["risks"]) else ""
        cells[1].text = ATAM_ANALYSIS["non_risks"][idx] if idx < len(ATAM_ANALYSIS["non_risks"]) else ""

    doc.add_heading("5 最终选型与理由", level=1)
    doc.add_paragraph(ATAM_ANALYSIS["recommendation"])
    doc.add_paragraph("最终建议采用“微服务主线 + 核心双服务先落地”的折中路线：实验二在逻辑上完整表达微服务边界，实验三实现上只优先落地 member-service 与 reservation-service，兼顾课程主题、架构可评审性和实现复杂度。")

    doc.add_heading("6 下一阶段输入", level=1)
    next_inputs = [
        "D2-4 将基于本报告的最终推荐方案展开 C4 多视图。",
        "D2-5 将按 member-service 与 reservation-service 先行设计 OpenAPI 契约。",
        "D2-6 将围绕六个上下文绘制 ER 图，并重点展开双核心服务的数据模型。",
    ]
    for item in next_inputs:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7 参考文献", level=1)
    for item in REFERENCES:
        doc.add_paragraph(item)

    doc.save(SOURCE_DIR / artifact_filename("D2-3", "docx"))


if __name__ == "__main__":
    ensure_dirs()
    build_d2_1()
    build_d2_2()
    build_d2_3()
